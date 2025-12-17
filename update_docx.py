from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn

# === CONFIGURATION ===
input_docx = 'Smart_Face_Sticker_Generator_Project_Report_Updated.docx'


output_docx = "Smart_Face_Sticker_Generator_Final_Report.docx"

# Images and their captions
images = [
    ("ORIGINALL IMG.jpg", "Figure 1: Original Input Image"),
    ("Edge Detection(Canny).jpg", "Figure 2: Edge Detection Output (Canny Edge Detector)"),
    ("Closed Edges.jpg", "Figure 3: Closed Edges after Morphological Operations"),
    ("Final Sticker.jpg", "Figure 4: Final Normal Sticker Output"),
    ("Sticker(Black and White).jpg", "Figure 5: Final Black and White Sticker Output"),
    ("Screenshot 2025-11-02 201029.png", "Figure 6: Streamlit Interface – Upload Section"),
    ("Screenshot 2025-11-02 202825.png", "Figure 7: Streamlit Interface – Sticker Statistics and Processing Steps")
]

# === Helper Functions ===
def add_separator_line(doc):
    """Adds a thin grey separator line between figures"""
    p = doc.add_paragraph()
    run = p.add_run("─" * 60)
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.alignment = 1

def set_page_numbers(doc):
    """Add page numbers to footer, starting from Abstract page"""
    sections = doc.sections
    for i, section in enumerate(sections):
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = 1
        # Add field for page number
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

# === MAIN SCRIPT ===
print("Loading document...")
doc = Document(input_docx)

# Find "5. OUTPUT IMAGES" section
output_index = None
for i, p in enumerate(doc.paragraphs):
    if "5. OUTPUT IMAGES" in p.text.upper():
        output_index = i
        break

if output_index is None:
    raise ValueError("Could not find '5. OUTPUT IMAGES' section in document.")

# Insert formal intro paragraph
intro = (
    "The following figures illustrate the step-by-step experimental results obtained during "
    "the implementation of the Smart Face Sticker Generator. Each figure represents a distinct "
    "phase of the image processing workflow — from preprocessing and edge detection to mask "
    "generation, sticker creation, and interface visualization. All results were generated using "
    "the Streamlit-based web interface developed for this project."
)

doc.paragraphs[output_index + 1].text = intro

# Insert images with captions and separators
for img_path, caption in images:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(img_path, width=Inches(5.5))
    p.alignment = 1  # center align image
    caption_p = doc.add_paragraph(caption)
    caption_p.style = "Normal"
    caption_p.alignment = 1
    run = caption_p.runs[0]
    run.bold = True
    add_separator_line(doc)

# Add page numbers
set_page_numbers(doc)

# Save the final version
doc.save(output_docx)
print(f"✅ Final project report generated successfully: {output_docx}")

